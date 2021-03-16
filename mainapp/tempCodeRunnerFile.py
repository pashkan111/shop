from docx import Document
doc = Document()
doc.add_heading('The title')
p = doc.add_paragraph('New paragraph')
doc.save('New_doc.doc')